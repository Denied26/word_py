from docxtpl import DocxTemplate
from docx import Document
from docx.enum.style import WD_STYLE_TYPE


# класс записки
class Note(object):
    # можно инициализировать отдельно все поля или передать массив
    # массив tabble_arr отвечает за таблицу, эго можно передать сюда или в функции добавления таблицы
    def __init__(self, to_position='', to_fio='', date='', number='', link='', theme='', to_io='', text='',
                 from_fio='', from_position='', note_arr='', performer='', t_number='', table_arr=None):
        if note_arr is not None:
            self.note_arr = note_arr
        else:
            self.note_arr = {
                'to_position': to_position,
                'to_fio': to_fio,
                'date': date,
                'number': number,
                'link': link,
                'theme': theme,
                'to_io': to_io,
                'text': text,
                'from_fio': from_fio,
                'from_position': from_position,
                'performer': performer,
                't_number': t_number
            }
        self.table_arr = table_arr
        # исходный ворд файл с перемеными
        self.source_file = "Служебная записка.docx"
        # если хотим сохранять не в папке проекта то добавть путь в инит или с помощью функции изменить
        self.save_place = ""
        # файл будет сохранятся под номером записки
        self.name_file = str(self.note_arr['number']) + '.docx'

    def chengeSavePlace(self, str):
        self.save_place = str

    # если передать ноль, то таблица не будет срендерена
    def makeWordFile(self, do_table=True):
        doc = DocxTemplate(self.source_file)
        doc.render(self.note_arr)
        doc.save(self.save_place + self.name_file)
        if do_table and self.table_arr is not None:
            self.addTableHard(self.table_arr)
    # рендер таблицы
    def formingTableArray(self, arr):
        #максимум столбцов с троке
        max_col = 0
        #строка где этот максимум
        row_max_col = 0
        #список столбцов
        col = {}
        # массив заполненый пустыми строками для заполнения в количестве столюцов
        empty_arr = []
        for i in range(len(arr)):
            if max_col < len(arr[i]):
                max_col = len(arr[i])
                row_max_col = i
        #готовый отсортированнный масссив с проставленными пустыми строками
        r_arr = []
        # формируем словарь ключ-индекс столбца
        for i, key in enumerate(arr[row_max_col]):
            col[key] = i
            empty_arr.append('')

        for row in range(len(arr)):
            r_arr.append([])
            r_arr[row] = empty_arr.copy()
            for key,value in arr[row].items():
                try :
                    r_arr[row][col[key]] = value
                except:
                    col[key] = len(col)
                    empty_arr.append('')
                    r_arr[row].append(value)
        return r_arr
    def addTableEasy(self,arr):
        doc = Document(self.name_file)
        text = doc.tables[0].rows[7].cells[1]
        table = text.add_table(rows=len(arr), cols=len(arr[0]))
        # границы таблицы
        table.style = 'Table Grid'
        for i in range(len(arr)):
            row = table.rows[i].cells
            for j in range(len(arr[0])):
                row[j].text = str(arr[i][j])
        doc.save(self.save_place + self.name_file)

    def addTableHard(self,t_arr):
        doc = Document(self.name_file)
        text = doc.tables[0].rows[7].cells[1]
        arr = self.formingTableArray(t_arr)
        max_col = 0
        # ищем строку с максимальным количеством столюцов
        for i in range(len(arr)):
            if len(arr[i]) > max_col:
                max_col = len(arr[i])
        table = text.add_table(rows=len(arr), cols=max_col)
        #границы таблицы
        table.style = 'Table Grid'
        # идем по строкам
        for i in range(len(arr)):
            row = table.rows[i].cells
            # идем по столбцам
            for j in range(len(arr[i])):
                row[j].text = str(arr[i][j])
                # если не дошел до макс кол столбцов забиваем пустой строкой
                for j in range(len(arr[i]), max_col):
                    row[j].text = ''
        doc.save(self.save_place + self.name_file)

    def addText(self,string):
        doc = Document(self.name_file)
        text = doc.tables[0].rows[7].cells[1]
        table = text.add_paragraph(string)
        doc.save(self.save_place + self.name_file)
